from flask import Flask, request, jsonify
from flask_cors import CORS
import PyPDF2
import docx
import re
from datetime import datetime
import json

app = Flask(__name__)
CORS(app)

# --- Expanded Skill List with categories ---
SKILLS = {
    "programming": ["python", "java", "c++", "c#", "javascript", "typescript", "go", "rust", "scala", "r"],
    "data_science": ["machine learning", "deep learning", "nlp", "computer vision", "data analysis", "statistics", "pandas", "numpy", "scikit-learn", "tensorflow", "pytorch", "keras"],
    "big_data": ["spark", "kafka", "hadoop", "pyspark", "hive", "storm", "flink", "elasticsearch", "mongodb", "cassandra"],
    "databases": ["sql", "mysql", "postgresql", "oracle", "sqlite", "nosql", "redis", "dynamodb"],
    "cloud": ["aws", "azure", "gcp", "docker", "kubernetes", "terraform", "jenkins", "ci/cd"],
    "web": ["html", "css", "react", "angular", "vue", "node.js", "express", "django", "flask", "spring"],
    "tools": ["git", "tableau", "power bi", "excel", "jira", "confluence", "linux", "unix"],
    "soft_skills": ["communication", "leadership", "teamwork", "problem solving", "analytical thinking", "project management"]
}

# Flatten skills for easy searching
ALL_SKILLS = []
for category, skills in SKILLS.items():
    ALL_SKILLS.extend(skills)

# --- Helper functions ---

def extract_text(file):
    """Enhanced text extraction from PDF or DOCX"""
    text = ""
    filename = file.filename.lower()
    
    try:
        if filename.endswith(".pdf"):
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + " "
        elif filename.endswith(".docx"):
            doc = docx.Document(file)
            for para in doc.paragraphs:
                text += para.text + " "
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
        else:
            # Handle plain text files
            text = file.read().decode("utf-8", errors='ignore')
            
        # Clean and normalize text
        text = re.sub(r'\s+', ' ', text)  # Replace multiple spaces with single space
        text = re.sub(r'[^\w\s]', ' ', text)  # Replace special chars with space
        return text.lower().strip()
        
    except Exception as e:
        print(f"Error extracting text from {filename}: {e}")
        return ""

def extract_candidate_name_from_text(resume_text):
    """Extract candidate name from resume text - optimized version"""
    if not resume_text:
        return None
    
    lines = resume_text.split('\n')
    # Check only first 3 lines for performance
    first_lines = [line.strip() for line in lines[:3] if line.strip()]
    
    # Simplified patterns for faster processing
    name_pattern = r'^([A-Z][a-z]+(?: [A-Z][a-z]+){1,2})'
    
    for line in first_lines:
        # Skip obvious headers
        if any(word in line.lower() for word in ['resume', 'cv', 'curriculum', '@', 'phone']):
            continue
            
        match = re.search(name_pattern, line)
        if match:
            name = match.group(1).strip().title()
            words = name.split()
            if 2 <= len(words) <= 3 and 4 <= len(name) <= 30:
                return name
    
    return None

def get_skill_suggestions(skill, category):
    """Generate specific suggestions based on skill and category"""
    suggestions = {
        "programming": f"Complete online courses for {skill} programming. Build 2-3 projects showcasing {skill} skills. Contribute to open-source {skill} projects.",
        "data_science": f"Take specialized courses in {skill}. Work on real-world datasets using {skill}. Create a portfolio project demonstrating {skill} expertise.",
        "big_data": f"Get hands-on experience with {skill} through cloud platforms. Complete certification courses. Build end-to-end projects using {skill}.",
        "databases": f"Practice {skill} queries and database design. Complete database administration courses. Build applications with {skill} integration.",
        "cloud": f"Pursue {skill} certifications. Set up personal projects using {skill} services. Learn infrastructure as code with {skill}.",
        "web": f"Build responsive web applications using {skill}. Complete full-stack development courses. Deploy projects showcasing {skill} skills.",
        "tools": f"Get proficient in {skill} through daily practice. Complete relevant certifications. Use {skill} in your projects and document the process.",
        "soft_skills": f"Develop {skill} through practical experience. Join relevant workshops or courses. Demonstrate {skill} through project leadership and team collaboration."
    }
    
    return suggestions.get(category, f"Learn {skill} through courses, tutorials, and hands-on practice. Add relevant projects to your portfolio.")

def calculate_score(resume_text, jd_text):
    """Enhanced scoring with better skill matching and categorized suggestions"""
    
    # Extract skills mentioned in JD
    jd_skills = []
    skill_categories = {}
    
    for category, skills in SKILLS.items():
        for skill in skills:
            if skill in jd_text.lower():
                jd_skills.append(skill)
                skill_categories[skill] = category
    
    # Find matched skills in resume
    matched_skills = []
    for skill in jd_skills:
        if skill in resume_text.lower():
            matched_skills.append(skill)
    
    missing_skills = list(set(jd_skills) - set(matched_skills))
    
    # Calculate score
    score = int((len(matched_skills) / len(jd_skills) * 100) if jd_skills else 0)
    
    # Determine verdict
    if score >= 80:
        verdict = "Excellent"
    elif score >= 60:
        verdict = "Good"
    elif score >= 40:
        verdict = "Average"
    else:
        verdict = "Needs Improvement"
    
    # Generate categorized improvement plan
    improvement_plan = []
    for skill in missing_skills:
        category = skill_categories.get(skill, "general")
        suggestion = get_skill_suggestions(skill, category)
        
        improvement_plan.append({
            "skill": skill.title(),
            "category": category.replace("_", " ").title(),
            "suggestion": suggestion,
            "priority": "High" if category in ["programming", "data_science"] else "Medium"
        })
    
    return score, verdict, missing_skills, matched_skills, improvement_plan

# --- Routes ---

@app.route("/match", methods=["POST"])
def match_resume():
    if "resume" not in request.files or "jd" not in request.files:
        return jsonify({"error": "Resume or JD file missing"}), 400
    
    resume_file = request.files["resume"]
    jd_file = request.files["jd"]
    
    if resume_file.filename == "" or jd_file.filename == "":
        return jsonify({"error": "No files selected"}), 400
    
    try:
        # Extract text once for both operations
        resume_text = extract_text(resume_file)
        jd_text = extract_text(jd_file)
        
        if not resume_text or not jd_text:
            return jsonify({"error": "Could not extract text from files"}), 400
        
        # Extract candidate name from resume text
        candidate_name = extract_candidate_name_from_text(resume_text)
        
        # Calculate matching score
        score, verdict, missing_skills, matched_skills, improvement_plan = calculate_score(resume_text, jd_text)
        
        # Prepare result with candidate name included
        result = {
            "score": score,
            "verdict": verdict,
            "missing_skills": missing_skills,
            "matched_skills": matched_skills,
            "improvement_plan": improvement_plan,
            "total_skills_required": len(missing_skills) + len(matched_skills),
            "skills_matched": len(matched_skills),
            "resume_filename": resume_file.filename,
            "jd_filename": jd_file.filename,
            "resume_word_count": len(resume_text.split()),
            "jd_word_count": len(jd_text.split()),
            "candidate_name": candidate_name,
            "timestamp": datetime.now().isoformat()
        }
        
        return jsonify(result)
        
    except Exception as e:
        print(f"Error processing files: {e}")
        return jsonify({"error": f"Error processing files: {str(e)}"}), 500

@app.route("/health", methods=["GET"])
def health_check():
    return jsonify({"status": "healthy", "timestamp": datetime.now().isoformat()})

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)