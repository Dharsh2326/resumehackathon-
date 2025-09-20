import PyPDF2
import docx
import re
import io
from typing import Optional, Tuple

def extract_text(file) -> str:
    """
    Enhanced text extraction that handles various file formats and encodings
    """
    text = ""
    filename = file.name.lower() if hasattr(file, 'name') else str(file).lower()
    
    try:
        if filename.endswith(".pdf"):
            text = extract_pdf_text(file)
        elif filename.endswith(".docx"):
            text = extract_docx_text(file)
        elif filename.endswith((".txt", ".doc")):
            text = extract_plain_text(file)
        else:
            # Try to read as plain text
            text = extract_plain_text(file)
            
        return clean_text(text)
        
    except Exception as e:
        print(f"Error extracting text from {filename}: {e}")
        return ""

def extract_pdf_text(file) -> str:
    """Extract text from PDF with better error handling"""
    text = ""
    try:
        # Handle both file path and file object
        if isinstance(file, str):
            with open(file, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        else:
            # File object
            file.seek(0)  # Reset file pointer
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                except Exception as e:
                    print(f"Error extracting page: {e}")
                    continue
                    
    except Exception as e:
        print(f"Error reading PDF: {e}")
        # Fallback: try to read as binary and decode
        try:
            file.seek(0)
            content = file.read()
            text = content.decode('utf-8', errors='ignore')
        except:
            pass
            
    return text

def extract_docx_text(file) -> str:
    """Extract text from DOCX including tables"""
    text = ""
    try:
        if isinstance(file, str):
            doc = docx.Document(file)
        else:
            file.seek(0)
            doc = docx.Document(file)
            
        # Extract paragraph text
        for para in doc.paragraphs:
            text += para.text + "\n"
            
        # Extract table text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
                
        # Extract headers and footers
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    text += para.text + "\n"
            if section.footer:
                for para in section.footer.paragraphs:
                    text += para.text + "\n"
                    
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        
    return text

def extract_plain_text(file) -> str:
    """Extract plain text with encoding detection"""
    text = ""
    encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
    
    try:
        if isinstance(file, str):
            with open(file, 'rb') as f:
                content = f.read()
        else:
            file.seek(0)
            content = file.read()
            
        # Try different encodings
        for encoding in encodings:
            try:
                if isinstance(content, bytes):
                    text = content.decode(encoding)
                else:
                    text = str(content)
                break
            except UnicodeDecodeError:
                continue
                
        if not text:
            # Last resort: ignore errors
            if isinstance(content, bytes):
                text = content.decode('utf-8', errors='ignore')
            else:
                text = str(content)
                
    except Exception as e:
        print(f"Error reading plain text: {e}")
        
    return text

def clean_text(text: str) -> str:
    """
    Enhanced text cleaning with better normalization
    """
    if not text:
        return ""
        
    # Convert to lowercase
    text = text.lower()
    
    # Remove extra whitespace and special characters
    text = re.sub(r'\s+', ' ', text)  # Multiple spaces to single space
    text = re.sub(r'[^\w\s\.\-\+]', ' ', text)  # Keep alphanumeric, spaces, dots, hyphens, plus
    
    # Remove standalone numbers and very short words (except common tech terms)
    words = text.split()
    tech_terms = {'c', 'r', 'ai', 'ml', 'ui', 'ux', 'qa', 'ci', 'cd', 'it', 'bi', 'etl'}
    filtered_words = []
    
    for word in words:
        if len(word) >= 2 or word in tech_terms:
            filtered_words.append(word)
    
    text = ' '.join(filtered_words)
    
    # Normalize common variations
    replacements = {
        'machine learning': 'machine_learning',
        'data science': 'data_science',
        'computer vision': 'computer_vision',
        'natural language processing': 'nlp',
        'artificial intelligence': 'ai',
        'deep learning': 'deep_learning',
        'big data': 'big_data',
        'cloud computing': 'cloud_computing',
        'software engineering': 'software_engineering',
        'web development': 'web_development',
        'mobile development': 'mobile_development',
        'database management': 'database_management',
        'project management': 'project_management',
        'quality assurance': 'qa',
        'user experience': 'ux',
        'user interface': 'ui'
    }
    
    for phrase, replacement in replacements.items():
        text = text.replace(phrase, replacement)
    
    return text.strip()

def extract_sections(text: str) -> dict:
    """
    Extract different sections from resume text
    """
    sections = {
        'education': '',
        'experience': '',
        'skills': '',
        'projects': '',
        'certifications': ''
    }
    
    # Common section headers
    section_patterns = {
        'education': r'(education|academic|qualification|degree)',
        'experience': r'(experience|employment|work|career|professional)',
        'skills': r'(skills|technical|competenc|proficienc)',
        'projects': r'(projects|portfolio|work)',
        'certifications': r'(certification|certificate|license|award)'
    }
    
    text_lower = text.lower()
    
    for section, pattern in section_patterns.items():
        matches = re.finditer(pattern, text_lower)
        for match in matches:
            start = match.start()
            # Try to find the end of this section (next section or end of text)
            end = len(text)
            for other_pattern in section_patterns.values():
                if other_pattern != pattern:
                    next_matches = re.finditer(other_pattern, text_lower[start+100:])
                    for next_match in next_matches:
                        potential_end = start + 100 + next_match.start()
                        if potential_end < end:
                            end = potential_end
                        break
            
            section_text = text[start:end]
            if len(section_text) > len(sections[section]):
                sections[section] = section_text
    
    return sections

def validate_file_format(file) -> Tuple[bool, str]:
    """
    Validate if the uploaded file is in supported format
    """
    if not hasattr(file, 'name'):
        return False, "Invalid file object"
    
    filename = file.name.lower()
    supported_extensions = ['.pdf', '.docx', '.txt', '.doc']
    
    if not any(filename.endswith(ext) for ext in supported_extensions):
        return False, f"Unsupported file format. Supported formats: {', '.join(supported_extensions)}"
    
    # Check file size (max 10MB)
    try:
        file.seek(0, 2)  # Seek to end
        size = file.tell()
        file.seek(0)  # Reset to beginning
        
        if size > 10 * 1024 * 1024:  # 10MB
            return False, "File size too large (max 10MB)"
            
    except Exception as e:
        print(f"Could not check file size: {e}")
    
    return True, "Valid file format"

def preprocess_text_for_matching(text: str) -> str:
    """
    Preprocess text specifically for skill matching
    """
    text = clean_text(text)
    
    # Expand common abbreviations
    abbreviations = {
        'js': 'javascript',
        'ts': 'typescript',
        'py': 'python',
        'db': 'database',
        'ai/ml': 'artificial intelligence machine learning',
        'nlp': 'natural language processing',
        'cv': 'computer vision',
        'dl': 'deep learning',
        'rnn': 'recurrent neural network',
        'cnn': 'convolutional neural network',
        'api': 'application programming interface',
        'rest': 'representational state transfer',
        'crud': 'create read update delete',
        'orm': 'object relational mapping',
        'mvc': 'model view controller',
        'spa': 'single page application',
        'pwa': 'progressive web application'
    }
    
    words = text.split()
    expanded_words = []
    
    for word in words:
        if word in abbreviations:
            expanded_words.extend(abbreviations[word].split())
        else:
            expanded_words.append(word)
    
    return ' '.join(expanded_words)